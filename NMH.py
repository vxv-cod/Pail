from PyQt5 import QtCore, QtGui, QtWidgets
import sys
import os
import pickle
from okno_general import ui
from okno_general import Form
from okno_general import app

# import docx
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
document = Document()

# Рисунок
from PyQt5.QtWidgets import QWidget
from PyQt5.QtWidgets import QApplication
from PyQt5.QtGui import QPainter
from PyQt5.QtGui import QBrush
from PyQt5.QtGui import QPen
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QPixmap

# import math
from math import cos
from math import tan

import vxv_tnnc_SQL_dll

_translate = QtCore.QCoreApplication.translate
Form.setWindowTitle(_translate("Form", "Pail"))
ui.tabWidget.setCurrentIndex(1)

# =============================================================================
# -----------------------------------------------------------------------------
# Дополняем таблицу Характеристики грунтов списком типов грунтов
ui.comboBox_0 = QtWidgets.QComboBox(ui.tab_2)
ui.comboBox_1 = QtWidgets.QComboBox(ui.tab_2)
ui.comboBox_2 = QtWidgets.QComboBox(ui.tab_2)
ui.comboBox_3 = QtWidgets.QComboBox(ui.tab_2)
ui.comboBox_4 = QtWidgets.QComboBox(ui.tab_2)
ui.comboBox_5 = QtWidgets.QComboBox(ui.tab_2)
ui.comboBox_6 = QtWidgets.QComboBox(ui.tab_2)
ui.comboBox_7 = QtWidgets.QComboBox(ui.tab_2)
ui.comboBox_8 = QtWidgets.QComboBox(ui.tab_2)
ui.comboBox_9 = QtWidgets.QComboBox(ui.tab_2)

for i in range(0, 10):
    eval('ui.comboBox_{}.setObjectName("comboBox_{}")'.format(i, i))
    eval('ui.comboBox_{}.addItem("")'.format(i))
    eval('ui.comboBox_{}.addItem("")'.format(i))
    eval('ui.comboBox_{}.addItem("")'.format(i))
    eval('ui.comboBox_{}.addItem("")'.format(i))
    eval('ui.comboBox_{}.addItem("")'.format(i))
    eval('ui.comboBox_{}.addItem("")'.format(i))
    eval('ui.comboBox_{}.addItem("")'.format(i))
    eval('ui.comboBox_{}.addItem("")'.format(i))
    eval('ui.comboBox_{}.addItem("")'.format(i))
    eval('ui.comboBox_{}.setItemText(0, _translate("Form", ""))'.format(i))
    eval('ui.comboBox_{}.setItemText(1, _translate("Form", "  глина"))'.format(i))
    eval('ui.comboBox_{}.setItemText(2, _translate("Form", "  суглинок"))'.format(i))
    eval('ui.comboBox_{}.setItemText(3, _translate("Form", "  супесь"))'.format(i))
    eval('ui.comboBox_{}.setItemText(4, _translate("Form", "  песок гравел."))'.format(i))
    eval('ui.comboBox_{}.setItemText(5, _translate("Form", "  песок крупный"))'.format(i))
    eval('ui.comboBox_{}.setItemText(6, _translate("Form", "  песок средний"))'.format(i))
    eval('ui.comboBox_{}.setItemText(7, _translate("Form", "  песок мелкий"))'.format(i))
    eval('ui.comboBox_{}.setItemText(8, _translate("Form", "  песок пылев."))'.format(i))
    eval('ui.comboBox_{}.setFrame(False)'.format(i))
    eval('ui.comboBox_{}.setStyleSheet("background-color: rgb(254, 254, 254);")'.format(i))
    eval('ui.tableWidget_3.setCellWidget({}, 6, ui.comboBox_{})'.format(i, i))

    # ui.comboBox_2.setCurrentIndex(3)
    # ui.comboBox_3.setCurrentIndex(3)
    # ui.comboBox_4.setCurrentIndex(2)
    # ui.comboBox_5.setCurrentIndex(2)
    # ui.comboBox_6.setCurrentIndex(2)
    # ui.comboBox_7.setCurrentIndex(7)

# -----------------------------------------------------------------------------
def text_centr(x):
    ui.textEdit_2.append('')
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 100, 150)) # цвет текста
    # ui.textEdit_2.setFontItalic(True) # курсивный текст
    ui.textEdit_2.setFontWeight(100) # жирный текст
    ui.textEdit_2.append(x)
    ui.textEdit_2.setAlignment(QtCore.Qt.AlignHCenter) # центруем текст внутри абзаца
    ui.textEdit_2.setFontWeight(1) # убираем жирный текст
    # ui.textEdit_2.setFontItalic(False) # убираем курсивный текст
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 0, 0))
    global document
    # document.add_paragraph(x)
    paragraph = document.add_paragraph()
    # paragraph.add_run(x).bold = True
    paragraph.add_run(x, style='Intense Emphasis').bold = True
    # paragraph.add_run(x).font.color.rgb = RGBColor(0, 100, 150)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

def text_centr_black(x):
    ui.textEdit_2.append(x)
    ui.textEdit_2.setAlignment(QtCore.Qt.AlignHCenter) # центруем текст внутри абзаца
    global document
    paragraph = document.add_paragraph(x)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
def text_abzac(x):
    ui.textEdit_2.append('       {}'.format(x))
    ui.textEdit_2.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
    global document
    document.add_paragraph(x)

def text_abzac_color(x):
    # ui.textEdit_2.append('')
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 100, 150))
    ui.textEdit_2.setFontWeight(100) # жирный текст
    ui.textEdit_2.append('       {}'.format(x))
    ui.textEdit_2.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 0, 0))
    ui.textEdit_2.setFontWeight(1) # жирный текст
    ui.textEdit_2.append('')
    global document
    paragraph = document.add_paragraph()
    paragraph.add_run(x, style='Intense Emphasis').bold = True

def text_abzac_color_111(x):
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 100, 150))
    ui.textEdit_2.setFontWeight(100) # жирный текст
    ui.textEdit_2.append('       {}'.format(x))
    ui.textEdit_2.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 0, 0))
    ui.textEdit_2.setFontWeight(1) # жирный текст
    global document
    paragraph = document.add_paragraph()
    paragraph.add_run(x, style='Intense Emphasis').bold = True
# -----------------------------------------------------------------------------
# Проверка ячеек на пустые значения и ","
def vvod(nomerwidgeta, strok, stolbec):
    xox = []; yoy = []
    for i in range(1, strok+1):
        x = eval('ui.tableWidget{}.item({}-1, {}).text()'.format(nomerwidgeta, i, stolbec))
        if x != '': 
            x = x.replace(',', '.')
        else: 
            x = 0
        xox.append(x)
        if '.' in str(x):
            yoy.append(float(x))
        else:
            try:
                yoy.append(int(x))
            except:
                yoy.append(str(x))
    return xox, yoy

def sbor_dannih():
    global dannie, ige_skv, nni, ige_xap, Jl_xap, e_xap, γ1_xap, С1_xap, fi1_xap, tip_grunta_index, tip_grunta
    ige_xap = vvod('_3', 10, 0)[0]
    Jl_xap = vvod('_3', 10, 1)[1]
    e_xap = vvod('_3', 10, 2)[1]
    γ1_xap = vvod('_3', 10, 3)[1]
    С1_xap = vvod('_3', 10, 4)[1]

    fi1_xap = vvod('_3', 10, 5)[1]
    tip_grunta = []
    tip_grunta_index = []
    for i in range(0, 10):
        x = eval('ui.comboBox_{}.currentText()'.format(i))
        y = eval('ui.comboBox_{}.currentIndex()'.format(i))
        if x != '': 
            tip_grunta.append(x)
            tip_grunta_index.append(y)
        else: 
            x = '0'
            tip_grunta.append(x)
            tip_grunta_index.append(y)
    dannie = vvod('_7', 14, 0)[1]        # исходные данные
    ige_skv = vvod('_5', 18, 0)[0]      # список ИГЭ скважины
    nni = vvod('_5', 18, 1)[1]          # список толщин слоев скважины

def delet_0(x):
    while x[-1] == 0: del x[-1]

# Вывод ошибки

def error_show(x):
    ui.label_26.setText(_translate("Form", "Ошибка"))
    ui.textEdit_2.setTextColor(QtGui.QColor (255, 0, 0))
    ui.textEdit_2.setText('Ошибка данных:')
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 0, 0))
    ui.textEdit_2.append(x)

def raschet():
    '''Отправляем сигнал на сервер об нажатии кнопки "Расчет" '''
    vxv_tnnc_SQL_dll.SendText('Pail_NMH')

    ui.textEdit_2.setText('')
    _translate = QtCore.QCoreApplication.translate
    ui.label_26.setText(_translate("Form", "Отчет"))
    global document
    del document    # удаляем глобальную переменную и заменяем ее новой такой же, но пустой
    document = Document()

    sbor_dannih()

    # Проверки данных
    for i in range(0, len(ige_skv)):
        if ige_skv[i] not in ige_xap:
            error_show('ИГЭ скважины не соответствуют ИГЭ в таблице \" Характеристики грунтов \"')
            return
    for i in range(0, len(ige_skv)):
        if ige_skv[i] == 0 and nni[i] != 0:
            error_show('ИГЭ скважины не заполнены!!!')
            return
    if sum(nni) == 0:
        error_show('Не заполнены данные по скважине!!!')
        return
    for i in range(0, len(ige_xap)):
        if ige_xap[i] != 0 and tip_grunta[i] == 0:
            error_show('Заполните все типы грунтов в таблице \" Характеристики грунтов \"')
            return
    
    # def delet_0(x):
    #     while x[-1] == 0: del x[-1]

    delet_0(ige_skv)
    delet_0(nni)

    # Данные сваи
    lo = dannie[0]             # Отметка верха сваи, м
    Ls = dannie[1]             # Длина сваи, 
    l = round(Ls  -  lo, 3)               # Заглубление сваи в грунт, м
    Ds = round(dannie[2] * 0.001, 3)      # Диаметр сваи, м
    Tss = round(dannie[3] * 0.001, 3)     # Толщина стенки сваи, м
    
    # Проверки данных
    if Ds == 0:
        error_show('Диаметр сваи не корректен!!!')
        return
    if Tss == 0:
        error_show('Тошщина стенки сваи не корректна!!!')
        return

    # Нагрузки
    Ns = dannie[5]      # Сжимающая нагрузка на сваю, тс
    Hs = dannie[6]      # Момент на сваю, тс
    Ms = dannie[7]      # Горизонтальная нагрузка на сваю, тс

    #Модуль упругости материала сваи:
    E = round(206000000 * 0.101971621298, 3)

    # Момент инерции сечения сваи:
    I = round((3.14159265359 * (Ds**4  -  (Ds  -  Tss*2)**4)) / 64, 5)

    # условная ширина сваи:
    bp = Ds + 1 if Ds >= 0.8 else 1.5 * Ds + 0.5

    # коэффициент условий работы:
    yc = 1

    # Функция интерполяции 
    def interpoi (t1, t2, yy):
        for i in t1:
            if yy > i:
                continue
            else:
                break
        ia1 = t1.index(i); ia2 = t1.index(i)-1
        a1 = t1[ia1]; a2 = t1[ia2]; b1 = t2[ia1]; b2 = t2[ia2]
        a12 = a1  -  a2; b12 = b1  -  b2; a13 = yy  -  a2; x = a13 * b12 / a12; y = b2 + x
        return y

    t1_glina_suglinok = [0, 0.5, 0.75, 1]
    t2_glina_suglinok = [6000, 4000, 2350, 1350]
    t2_glina_suglinok = [round(i * 0.101971621298, 2) for i in t2_glina_suglinok]

    t1_supes = [0, 0.5, 0.75]
    t2_supes = [4000, 2888.9, 2350]
    t2_supes = [round(i * 0.101971621298, 4) for i in t2_supes]

    t1_pesok_grav = [0.55, 0.6, 0.65, 0.7]
    t2_pesok_grav = [33350, 27816.7, 22283.3, 16750]
    t2_pesok_grav = [round(i * 0.101971621298, 4) for i in t2_pesok_grav]

    t1_pesok_krup = [0.55, 0.6, 0.65, 0.7]
    t2_pesok_krup = [10000, 8666.7, 7333.3, 6000]
    t2_pesok_krup = [round(i * 0.101971621298, 4) for i in t2_pesok_krup]

    t1_pesok_sred = [0.55, 0.6, 0.65, 0.7]
    t2_pesok_sred = [6000, 5333.3, 4666.7, 4000]
    t2_pesok_sred = [round(i * 0.101971621298, 4) for i in t2_pesok_sred]

    t1_pesok_melk = [0.6, 0.65, 0.7, 0.75]
    t2_pesok_melk = [6000, 5333.3, 4666.7, 4000]
    t2_pesok_melk = [round(i * 0.101971621298, 4) for i in t2_pesok_melk]

    t1_pesok_pile = [0.6, 0.65, 0.7, 0.75, 0.8]
    t2_pesok_pile = [4000, 3583.3, 3166.7, 2750, 2350]
    t2_pesok_pile = [round(i * 0.101971621298, 4) for i in t2_pesok_pile]

    K_xap = []; pesok_glinistiu = []
    for i in range(0, len(ige_xap)):
        if tip_grunta[i] == '0':
            K_xap.append('0')
            pesok_glinistiu.append('0')
        if tip_grunta[i] == "  глина":
            t1 = t1_glina_suglinok
            t2 = t2_glina_suglinok
            K_xap.append(interpoi(t1, t2, Jl_xap[i]))
            pesok_glinistiu.append('1')
        if tip_grunta[i] == "  суглинок":
            t1 = t1_glina_suglinok
            t2 = t2_glina_suglinok
            K_xap.append(interpoi(t1, t2, Jl_xap[i]))
            pesok_glinistiu.append('1')
        if tip_grunta[i] == "  супесь":
            t1 = t1_supes
            t2 = t2_supes
            K_xap.append(interpoi(t1, t2, Jl_xap[i]))
            pesok_glinistiu.append('1')
        if tip_grunta[i] == "  песок гравел.":
            t1 = t1_pesok_grav
            t2 = t2_pesok_grav
            K_xap.append(interpoi(t1, t2, e_xap[i]))
            pesok_glinistiu.append('2')
        if tip_grunta[i] == "  песок крупный":
            t1 = t1_pesok_krup
            t2 = t2_pesok_krup
            K_xap.append(interpoi(t1, t2, e_xap[i]))
            pesok_glinistiu.append('2')
        if tip_grunta[i] == "  песок средний":
            t1 = t1_pesok_sred
            t2 = t2_pesok_sred
            K_xap.append(interpoi(t1, t2, e_xap[i]))
            pesok_glinistiu.append('2')
        if tip_grunta[i] == "  песок мелкий":
            t1 = t1_pesok_melk
            t2 = t2_pesok_melk
            K_xap.append(interpoi(t1, t2, e_xap[i]))
            pesok_glinistiu.append('2')
        if tip_grunta[i] == "  песок пылев.":
            t1 = t1_pesok_pile
            t2 = t2_pesok_pile
            K_xap.append(interpoi(t1, t2, e_xap[i]))
            pesok_glinistiu.append('2')
        
    # функция сбора списков для скважины из характеристик грунта
    def dannie_skv(x):      
        y = {}
        for i in range(0, len(ige_xap)):
            y[ige_xap[i]] = x[i]
        return [y[i] for i in ige_skv]

    K_skv = dannie_skv(K_xap)                   # список К для скважины
    fi1_skv = dannie_skv(fi1_xap)               # список угла внутреннего трения грунта скважины, град
    С1_skv = dannie_skv(С1_xap)                 # список удельного сцепления грунта скважины, тс/м
    γ1_skv = dannie_skv(γ1_xap)                 # список удельный (объемный) вес грунта скважины, тс/м3
    tip_grunta_skv = dannie_skv(tip_grunta)     # список типов грунта скважины
    pesok_glinistiu_skv = dannie_skv(pesok_glinistiu)     # список типов грунта скважины 0,1,2
    Jl_skv = dannie_skv(Jl_xap)                 # список Jl скважины
    e_skv = dannie_skv(e_xap)                   # список e скважины

    # Глубина в пределах которой определяется коэффициент К:
    lk = 3.5 * Ds + 1.5

    # приведенное значение коэффициента пропорциональности K
    if sum(nni) < lk:
        error_show('Недостаточная глубина по скважине!!!')
        return
    if l < lk:
        error_show('Длина сваи не корректна!!!')
        return

    if lk <= nni[0]:
        K = round(K_skv[0], 5)
    else:
        K = round(( K_skv[0] * nni[0] * (2 * lk  -  nni[0]) + K_skv[1] * (lk  -  nni[0])**2 ) / lk**2, 5)

    # Расчет коэффициента деформации αε 
    ae = round(((K * bp) / (yc * E * I))**0.2, 5)
    
    l1 = lo + 2 / ae; l1 = round(l1, 3)

    # приведенная глубины погружения сваи в грунт l' 
    lp = round(ae * l, 5)
    lp_D2 = round(ae * l, 1)
    if lp <= 2.0:
        lp_D2 = round(ae * l, 1)
    if lp > 2.0 and lp <= 3.0:
        lp_D2 = round(round(lp/0.2) * 0.2, 1)
    if lp > 3.0 and lp < 4.0:
        lp_D2 = round(round(lp/0.5) * 0.5, 1)

    # таблица Д2
    t1_lp = [0.5, 0.6, 0.7, 0.8, 0.9, 1.0, 1.1,	1.2, 1.3, 1.4, 1.5, 1.6, 1.7, 1.8, 1.9, 2.0, 2.2, 2.4, 2.6, 2.8, 3.0, 3.5, 4.0]
    t2_Ao = [72.004, 50.007, 36.745, 28.140, 22.244, 18.030, 14.916, 12.552, 10.717, 9.266, 8.101, 7.154, 6.375, 5.730, 5.190, 4.737, 4.032, 3.526, 3.163, 2.905, 2.727, 2.502, 2.441]
    t2_Bo = [192.026, 111.149, 70.023, 46.943, 33.008, 24.106, 18.160, 14.041, 11.103, 8.954, 7.349, 6.129, 5.189, 4.456, 3.878, 3.418, 2.756, 2.327, 2.048, 1.869, 1.758, 1.641, 1.621]
    t2_Co = [576.243, 278.069, 150.278, 88.279, 55.307, 36.486, 25.123, 17.944, 13.235, 10.050, 7.838, 6.268, 5.133, 4.299, 3.679, 3.213, 2.591, 2.227, 2.013, 1.889, 1.818, 1.757, 1.751]
    if lp >= t1_lp[-1]:
        Ao = t2_Ao[-1]
        Bo = t2_Bo[-1]
        Co = t2_Co[-1]
    else:
        t1 = t1_lp
        yy = lp_D2
        print('lp_D2 = ', lp_D2)
        t2 = t2_Ao
        Ao = round(interpoi(t1, t2, yy), 3)
        t2 = t2_Bo
        Bo = round(interpoi(t1, t2, yy), 3)
        t2 = t2_Co
        Co = round(interpoi(t1, t2, yy), 3)

    E = round(206000000 * 0.101971621298, 3)
    
    eHH = Ao / (ae**3 * E * I); eHH = round(eHH, 7)
    eMH = Bo / (ae**2 * E * I); eMH = round(eMH, 8)
    import copy
    eHM = copy.copy(eMH)
    eMM = Co / (ae * E * I); eMM = round(eMM, 8)

    # Горизонтальное перемещение Uo
    Ho = Hs; Mo = round(Ms + Hs * lo, 3)
    Uo = Ho * eHH + Mo * eHM; Uo = round(Uo, 5)
    wo = Ho * eMH + Mo * eMM; wo = round(wo, 5)

    # Расчетные значения горизонтального перемещения сваи в уровне подошвы ростверка Up, м, и угол ее поворота wp, рад
    Up = Uo + wo * lo + (Hs * lo**3) / (3 * E * I) + (Ms * lo**2) / (2 * E * I); Up = round(Up, 5)
    wp = wo + (Hs * lo**2)/(2 * E * I)  +  (Ms * lo) / (E * I); wp = round(wp, 5)

    # # Расчетная длина сваи на расчет гибкости и прогиба:
    # dcav = (Ds  -  Tss*2)
    # gibkost = l1 / ((I/(3.14159*((Ds/2)**2-(dcav/2)**2)))**0.5)
    
    # Предельный прогиб, м
    if dannie[9] == 0:
        Uu = round(l1/75, 3)
    else:
        Uu = round(float(dannie[9]), 3)
    if dannie[10] == 0:
        wu = round(float(dannie[10]), 3)
    else:
        wu = round(float(dannie[10]) * 0.0174533, 5)

    '''=============================================================================='''
    def GzMzQz(Z, Zp):
        # расчетные значения угла внутреннего трения грунта, радиан, и удельного сцепления грунта, кПа (тс/м );
        global y1, C1, fi1, ks, h1, h2, n, fM, fH, Mc, Mt
        global A1, B1, Ce1, D1, A3, B3, Ce3, D3, A4, B4, Ce4, D4, R, Gz, Mz, Qz
        if Z <= nni[0]: 
            fi1 = round(fi1_skv[0] / 57.2958, 4)
            C1 = round(С1_skv[0] * 0.101972, 5)
            y1 = γ1_skv[0]
        else:
            fi1 = round(fi1_skv[1] / 57.2958, 4)
            C1 = round(С1_skv[1] * 0.101972, 5)
            y1 = γ1_skv[1]
        # Проверка ограничения давления на грунт R:
        fM = dannie[12]     # отношение кратковременной нагрузки М к постоянной М
        fH = dannie[13]     # отношение кратковременной нагрузки H к постоянной H
        n = 4; h1 = 1; ks = 0.6
        Mc = round(Ms * (1  -  fM) + Hs * l1 * (1  -  fH), 3)
        Mt = round(Ms * fM  + Hs * l1 * fH, 3)
        h2 = round((Mc + Mt) / (n * Mc + Mt), 3)
        R = round(h1 * h2 * ( 4 / cos(fi1) ) * ( y1 * Z * tan(fi1) + ks * C1 ), 3)
        t2_Z = 0, 0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7, 0.8, 0.9, 1, 1.1, 1.2, 1.3, 1.4, 1.5, 1.6, 1.7, 1.8, 1.9, 2, 2.2, 2.4, 2.6, 2.8, 3, 3.5, 4
        t2_A1 = 1, 1, 1, 1, 1, 1, 0.999, 0.999, 0.997, 0.995, 0.992, 0.987, 0.979, 0.969, 0.955, 0.937, 0.913, 0.882, 0.843, 0.795, 0.735 , 0.575, 0.347, 0.033, -0.385, -0.928, -2.928, -5.853
        t2_B1 = 0,	0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7, 0.799, 0.899, 0.997, 1.095, 1.192, 1.287, 1.379, 1.468, 1.553, 1.633, 1.706, 1.77, 1.823, 1.887, 1.874, 1.755, 1.49, 1.037, -1.272, -5.941
        t2_C1 = 0,	0.005, 0.02, 0.045, 0.08, 0.125, 0.18, 0.245, 0.32, 0.405, 0.499, 0.604, 0.718, 0.841, 0.974, 1.115, 1.264, 1.421, 1.584, 1.752, 1.924, 2.272, 2.609, 2.907, 3.128, 3.225, 2.463, -0.927
        t2_D1 = 0,	0, 0.001, 0.005, 0.011, 0.021, 0.036, 0.057, 0.085, 0.121, 0.167, 0.222, 0.288, 0.365, 0.456, 0.56, 0.678, 0.812, 0.961, 1.126, 1.308, 1.72, 2.195, 2.724, 3.288, 3.858, 4.98, 4.548
        t2_A3 = 0,	0, -0.001, -0.005, -0.011, -0.021, -0.036, -0.057, -0.085, -0.121, -0.167, -0.222, -0.287, -0.365, -0.455, -0.559, -0.676, -0.808, -0.956, -1.118, -1.295, -1.693, -2.141, -2.621, -3.103, -3.541, -3.919, -1.614
        t2_B3 = 0,	0, 0, -0.001, -0.002, -0.005, -0.011, -0.02, -0.034, -0.055, -0.083, -0.122, -0.173, -0.238, -0.319, -0.42, -0.543, -0.691, -0.867, -1.074, -1.314, -1.906, -2.663, -3.6, -4.718, -6, -9.544, -11.731
        t2_C3 = 1,	1, 1, 1, 1, 0.999, 0.998, 0.996, 0.992, 0.985, 0.975, 0.96, 0.938, 0.907, 0.866, 0.811, 0.739, 0.646, 0.53, 0.385, 0.207, -0.271, -0.949, -1.877, -3.108, -4.688, -10.34, -17.919
        t2_D3 = 0,	0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.699, 0.799, 0.897, 0.994, 1.09, 1.183, 1.273, 1.358, 1.437, 1.507, 1.566, 1.612, 1.64, 1.646, 1.575, 1.352, 0.917, 0.197, -0.891, -5.854, -15.076
        t2_A4 = 0,	-0.005, -0.02, -0.045, -0.08, -0.125, -0.18, -0.245, -0.32, -0.404, -0.499, -0.603, -0.716, -0.838, -0.967, -1.105, -1.248, -1.396, -1.547, -1.699, -1.848, -2.125, -2.339, -2.437, -2.346, -1.969, 1.074, 9.244
        t2_B4 = 0,	0, -0.003, -0.009, -0.021, -0.042, -0.072, -0.114, -0.171, -0.243, -0.333, -0.443, -0.575, -0.73, -0.91, -1.116, -1.35, -1.613, -1.906, -2.227, -2.578, -3.36, -4.228, -5.14, -6.023, -6.765, -6.789, -0.358
        t2_C4 = 0,	0, 0, -0.001, -0.003, -0.008, -0.016, -0.03, -0.051, -0.082, -0.125, -0.183, -0.259, -0.356, -0.479, -0.63, -0.815, -1.036, -1.299, -1.608, -1.966, -2.849, -3.973, -5.355, -6.99, -8.84, -13.692, -15.611
        t2_D4 = 1,	1, 1, 1, 1, 0.999, 0.997, 0.994, 0.989, 0.98, 0.967, 0.946, 0.917, 0.876, 0.821, 0.747, 0.652, 0.529, 0.374, 0.181, -0.057, -0.692, -1.592, -2.821, -4.445, -6.52, -13.826, -23.14
        A1 = round(interpoi(t2_Z, t2_A1, Zp), 5)
        B1 = round(interpoi(t2_Z, t2_B1, Zp), 5)
        Ce1 = round(interpoi(t2_Z, t2_C1, Zp), 5)
        D1 = round(interpoi(t2_Z, t2_D1, Zp), 5)
        A3 = round(interpoi(t2_Z, t2_A3, Zp), 5)
        B3 = round(interpoi(t2_Z, t2_B3, Zp), 5)
        Ce3 = round(interpoi(t2_Z, t2_C3, Zp), 5)
        D3 = round(interpoi(t2_Z, t2_D3, Zp), 5)
        A4 = round(interpoi(t2_Z, t2_A4, Zp), 5)
        B4 = round(interpoi(t2_Z, t2_B4, Zp), 5)
        Ce4 = round(interpoi(t2_Z, t2_C4, Zp), 5)
        D4 = round(interpoi(t2_Z, t2_D4, Zp), 5)
        Gz = round(K * Zp / ae * (Uo * A1  -  wo * B1 / ae + Mo * Ce1 / (ae**2 * E * I) + Ho * D1 / ( ae**3 * E * I)), 3)
        Mz = round(ae**2 * E * I * Uo * A3 - ae * E * I * wo * B3 + Mo * Ce3 + Ho * D3 / ae, 2)
        Qz = round(ae**3 * E * I * Uo * A4 - ae**2 * E * I * wo * B4 + ae * Mo * Ce4 + Ho * D4, 2)
    
    # --------------------------------------------------------------------------
    ui.textEdit_2.setFontWeight(100)
    text_centr('''Расчет свай на совместное действие вертикальной и горизонтальной сил и момента.''')
    ui.textEdit_2.setFontWeight(1)

    ui.textEdit_2.setFontUnderline(True)
    text_centr('Исходные данные:')
    ui.textEdit_2.setFontUnderline(False)
    ui.textEdit_2.append('')

    text_abzac_color('Параметры сваи:')
    text_abzac('\tДлина: {} м'.format(Ls))
    text_abzac('\tОтметка верха: {} м'.format(lo))
    # if ui.radioButton.isChecked() == True:
    text_abzac('\tДиаметр: {} м'.format(Ds))
    text_abzac('\tТолщина стенки: {} м'.format(Tss))
    ui.textEdit_2.append('')

    text_abzac_color('Нагрузка на сваю:')
    text_abzac('\tСжимающая нагрузка на голову сваи N = {} тс'.format(round(Ns,3)))
    text_abzac('\tИзгибающий момент на голову сваи М = {} тсм'.format(round(Ms,3)))
    text_abzac('\tГоризонтальная нагрузка на голову сваи H = {} тс'.format(round(Hs,3)))
    ui.textEdit_2.append('')

    text_abzac_color_111('Параметры грунтов:')
    for i in range(0, len(tip_grunta_skv)):
        ui.textEdit_2.setFontWeight(100)
        ui.textEdit_2.append('')
        text_abzac_color('Слой №{}:'.format(i+1))
        ui.textEdit_2.setFontWeight(1)

        text_abzac('\tТип грунта  -{}'.format(tip_grunta_skv[i]))
        text_abzac('\tТолщина грунта h = {} м'.format(round(nni[i], 3)))
        text_abzac('\tУдельный вес грунта γI = {} тс/м3'.format(γ1_skv[i]))
        text_abzac('\tУгол внутреннего трения фI = {} град. = {} рад.'.format(fi1_skv[i], round(fi1_skv[i] / 57.2958, 4)))
        text_abzac('\tУдельное сцепление грунта С1 = {} кПа = {} тс/м2'.format(С1_skv[i], round(С1_skv[i] * 0.101972, 3)))
        if pesok_glinistiu_skv[i] == '1': text_abzac('\tПоказатель текучести грунта Jl = {}'.format(Jl_skv[i]))
        text_abzac('\tКоэффициент пористости e = {}'.format(e_skv[i]))
    # ==========================================================

    ui.textEdit_2.setFontUnderline(True)
    text_centr('Результаты расчета')
    ui.textEdit_2.setFontUnderline(False)
    ui.textEdit_2.append('')

    if lk > nni[0]:
        text_abzac('При наличии в пределах длины сваи нескольких слоев грунта рекомендуется для определения сопротивления грунта на боковой поверхности сваи пользоваться одним приведенным значением коэффициента пропорциональности K, принимаемым в зависимости от грунтов, расположенных до глубины lk, м, отсчитываемой от поверхности грунта при высоком ростверке или от подошвы ростверка при низком ростверке по СП 50-102-2003, где d  -  диаметр сваи:')
        text_centr('lk = 3.5 ∙ d + 1.5 м = 3.5 ∙ {} м + 1.5 м = {} м'.format(Ds, lk))
        ui.textEdit_2.append('')
        text_abzac('Если в пределах глубины lk расположено два слоя грунта, то приведенное значение K определяют по формуле:')
        text_centr('K = (KI ∙ lI ∙ (2 ∙ lk - lI) + KII ∙ (lk  -  lI)^2)/lk^2')
        text_abzac('где:')
        text_abzac('lI  -  толщина первого (верхнего) слоя грунта, м;')
        text_abzac('KI и KII  -  коэффициенты пропорциональности, принимаемые по таблице В.1, для грунтов I и II слоев, тс/м4.')
        text_centr('K = ({} тс/м4 ∙ {} м ∙ (2 ∙ {} м - {} м) + {} тс/м4 ∙ ({} м - {} м)^2)/{} м^2 = {} тс/м4'.format(K_skv[0], nni[0], lk, nni[0], K_skv[1], lk, nni[0], lk, K))
        ui.textEdit_2.append('')
    
    
    text_abzac('В.5    Расчеты по определению прочности свай всех видов следует проводить с учетом п. 7.1.8 с использованием коэффициента деформации αε (1/м), определяемого по формуле:')
    text_centr('αε = ((K ∙ bp)/(γc ∙ E ∙ I))^0.2')
    text_abzac('где:')
    text_abzac('γc = 1  -  коэффициент условий работы;')
    text_abzac('E = {} тс/м2  -  модуль упругости материала сваи;'.format(E))
    text_abzac('I = {} м4  -  момент инерции поперечного сечения сваи;'.format(I))
    if Ds < 0.8:
        text_abzac('bp = 1.5 ∙ {} м + 0.5 м = {} м  -  условная ширина сваи, м;'.format(Ds, bp))
    else:
        text_abzac('bp = {} м + 1.0 м = {} м  -  условная ширина сваи, м;'.format(Ds, bp))

    if lk > nni[0]:
        text_abzac('К = {} тс/м4  -  приведенное значение коэффициента пропорциональности.'.format(K))
    else:
        text_abzac('К = {} тс/м4  -  коэффициент пропорциональности, принимаемый в зависимости от вида грунта, окружающего сваю, по таблице В.1.'.format(K))
    text_centr('αε = (({} тс/м4 ∙ {} м)/({} ∙ {} тс/м2 ∙ {} м4))^0.2 = {} (1/м)'.format(K, bp, yc, E, I, ae))
    ui.textEdit_2.append('')

    text_abzac('7.1.8    При расчете свай всех видов по прочности материала сваю допускается рассматривать как стержень, жестко защемленный в грунте в сечении, расположенном от подошвы ростверка на расстоянии, определяемом по формуле:')
    text_centr('l1 = lo + 2/ae = {} м + 2/{} (1/м) = {} м'.format(lo, ae, l1))
    text_abzac('где:')
    text_abzac('lo = {} м - длина участка сваи от подошвы высокого ростверка до уровня планировки грунта, м;'.format(lo))
    text_abzac('αε = {} (1/м)  -  коэффициент деформации;'.format(ae))
    ui.textEdit_2.append('')

    text_abzac('Д.3    Расчеты свай по предельным состояниям двух групп следует выполнять с использованием значений приведенной глубины погружения сваи в грунт lp и приведенной глубины расположения сечения сваи в грунте  Zp, определяемых по формулам:')
    text_centr('lp = αε ∙ l = {} (1/м) ∙ {} м = {}'.format(ae, l, lp))
    text_abzac('где:')
    text_abzac('αε = {} (1/м)  -  коэффициент деформации;'.format(ae))
    text_abzac('l = {} м  -  действительная глубина погружения сваи (ее нижнего конца) в грунт.'.format(l))
    ui.textEdit_2.append('')
    
    text_abzac('Д.5.    Горизонтальное перемещение Uo, м, и угол поворота ψo, рад в уровне поверхности грунта при высоком ростверке, а при низком ростверке  -  в уровне его подошвы, следует определять по формулам:')
    text_centr('Uo = Ho ∙ εHH + Mo ∙ εHM')
    text_centr('ψo = Ho ∙ εMH + Mo ∙ εMM')
    text_abzac('где:')
    text_abzac('Ho = H = {} тс  -  расчетное значение поперечной силы в уровне земли;'.format(Hs))
    text_abzac('Mo = M + H ∙ lo = {} тсм + {} тс ∙ {} м = {} тсм  -  расчетное значение изгибающего момента в уровне земли, тсм, где H и M   -  расчетные значения поперечной силы и изгибающего момента, действующие на голову сваи;'.format(Ms, Hs, lo, Mo))
    text_abzac('Ao = {}; Bo = {}; Co = {}  -  безразмерные коэффициенты, принимаемые по таблице Д.2 в зависимости от lp;'.format(Ao, Bo, Co))
    
    # text_abzac('εHH = Ao / ( αε^3 ∙ E ∙ I )  -  горизонтальное перемещение сечения от силы Ho = 1 в уровне земли, м/тс:')
    # text_abzac('εMH = Bo / ( αε^2 ∙ E ∙ I )  -  угол поворота сечения от силы Ho = 1 от уровня земли, 1/тс;')
    # text_abzac('εHM = eMH  -  горизонтальное перемещение сечения от момента Mo = 1 в уровне земли, 1/(тсм);')
    # text_abzac('εMM = Co / ( αε ∙ E ∙ I )  -  угол поворота сечения от момента Mo = 1 от уровня земли, 1/(тсм);')
    # text_abzac('\tεHH = {} / ( {}^3 ∙ {} ∙ {} ) = {} м/тс'.format(Ao, ae, E, I, eHH))
    # text_abzac('\teHM = eMH = {} / ( {}^2 ∙ {} ∙ {} ) = {}'.format(Bo, ae, E, I, eMH))
    # text_abzac('\teMM = {} / ( {} ∙ {} ∙ {} ) = {}'.format(Co, ae, E, I, eMM))

    text_abzac('εHH = Ao/(αε^3 ∙ E ∙ I) = {} м/тс  -  горизонтальное перемещение сечения от силы Ho = 1 в уровне земли;'.format(eHH))
    text_abzac('εMH = Bo/(αε^2 ∙ E ∙ I) = {} (1/тс)  -  угол поворота сечения от силы Ho = 1 от уровня земли;'.format(eMH))
    text_abzac('εHM = eMH = {} (1/(тсм))  -  горизонтальное перемещение сечения от момента Mo = 1 в уровне земли;'.format(eHM))
    text_abzac('εMM = Co/(αε ∙ E ∙ I) = {} (1/(тсм))  -  угол поворота сечения от момента Mo = 1 от уровня земли.'.format(eMM))
    text_centr('Uo = {} тс ∙ {} м/тс + {} тсм ∙ {} (1/(тсм)) = {} м'.format(Ho, eHH, Mo, eHM, Uo))
    text_centr('ψo = {} тс ∙ {} (1/тс) + {} тсм ∙ {} (1/(тсм)) = {} рад'.format(Ho, eMH, Mo, eMM, wo))
    ui.textEdit_2.append('')

    text_abzac('Д.4.    Расчетные значения горизонтального перемещения сваи в уровне подошвы ростверка Up, м, и угол ее поворота Ψp, рад, следует определять по формулам:')
    text_centr('Up = Uo + ψo ∙ lo + (H ∙ lo^3)/(3 ∙ E ∙ I) + (M ∙ lo^2)/(2 ∙ E ∙ I)')
    text_centr('Ψp = ψo + (H ∙ lo^2)/(2 ∙ E ∙ I) + (M ∙ lo)/(E ∙ I)')
    text_abzac('где:')
    text_abzac('H = {} тс и M = {} тсм  -  расчетные значения поперечной силы и изгибающего момента, действующие на голову сваи;'.format(Hs, Ms))
    text_abzac('lo = {} м  -  длина участка сваи равная расстоянию от подошвы ростверка до поверхности грунта под ростверком, м;'.format(lo))
    text_abzac('E = {} тс/м2  -  модуль упругости материала сваи;'.format(E))
    text_abzac('I = {} м4  -  момент инерции поперечного сечения сваи;'.format(I))
    text_abzac('Uo и ψo  -  горизонтальное перемещение, м, и угол поворота поперечного сечения сваи, рад, в уровне поверхности грунта при высоком ростверке, а при низком ростверке  -  в уровне его подошвы, определяемые по Д.5.')
    text_centr('Up = {} м + {} ∙ {} м + ({} тс ∙ ({} м)^3)/(3 ∙ {} тс/м2 ∙ {} м4) + ({} тсм ∙ ({} м)^2)/(2 ∙ {} тс/м2 ∙ {} м4) = {} м'.format(Uo, wo, lo, Hs, lo, E, I, Ms, lo, E, I, Up))
    text_centr('Ψp = {} + ({} тс ∙ ({} м)^2)/(2 ∙ {} тс/м2 ∙ {} м4) + ({} тсм ∙ {} м)/({} тс/м2 ∙ {} м4) = {} рад'.format(wo, Hs, lo, E, I, Ms, lo, E, I, wp))
    ui.textEdit_2.append('')

    text_abzac('В.3    Расчет по предельному состоянию второй группы сводится к проверке соблюдения условий допустимости расчетных значений горизонтального перемещения голов свай и угла их поворота:')
    if Up <= Uu:
        ff = round(Up / Uu * 100, 1)
        text_centr('Up ≤ Uu  →  {} м ≤ {} м ({} % от предельного значения) - условие выполнено '.format(Up, Uu, ff))
    else:
        ui.textEdit_2.setTextColor(QtGui.QColor (255, 0, 0))
        text_centr_black('Up ≤ Uu  →  {} м ≥ {} м - условие НЕ выполнено'.format(Up, Uu))
        ui.textEdit_2.setTextColor(QtGui.QColor (0, 0, 0))
    if wp <= wu:
        ff = round(wp / wu * 100, 1)
        text_centr('Ψp ≤ Ψu  →  {} рад ≤ {} рад ({} % от предельного значения) - условие выполнено '.format(wp, wu, ff))
    else:
        ui.textEdit_2.append('')
        ui.textEdit_2.setTextColor(QtGui.QColor (255, 0, 0))
        text_centr_black('Ψp ≤ Ψu  →  {} рад ≥ {} рад - условие НЕ выполнено'.format(wp, wu))
        ui.textEdit_2.setTextColor(QtGui.QColor (0, 0, 0))

    text_abzac('где:')
    text_abzac('Up, Ψp  -  расчетные значения горизонтального перемещения сваи в уровне подошвы ростверка Up, м, и угол ее поворота Ψp, рад по Д.4;')
    if dannie[9] == 0:
        text_abzac('Uu = l1 / 75 = {} м / 75 = {} м  -  предельные допустимые значения горизонтального перемещения головы сваи, где l1 определяется по 7.1.8;'.format(l1, Uu))
    else:
        text_abzac('Uu = {} м  -  предельные допустимые значения горизонтального перемещения головы сваи;'.format(Uu))
    text_abzac('Ψu = {} рад  -  предельный угол поворота сваи.'.format(wu))
    ui.textEdit_2.append('')

    text_abzac('Д.7.   Расчетное давление Gz на грунт по контакту с боковой поверхностью сваи, возникающее на глубине Z, а также расчетный изгибающий момент Mz, поперечная сила Qz и продольная сила Nz, действующие на глубине Z в сечении сваи:')
    text_centr('Gz = (K/αε) ∙ Zp ∙ (Uo ∙ A1 - ψo ∙ B1/αε + Mo ∙ C1/(αε^2 ∙ E ∙ I) + Ho ∙ D1/(αε^3 ∙ E ∙ I))')
    text_centr('Mz = αε^2 ∙ E ∙ I ∙ Uo ∙ A3 - αε ∙ E ∙ I ∙ wo ∙ B3 + Mo ∙ C3 + Ho ∙ D3/ae')
    text_centr('Qz = αε^3 ∙ E ∙ I ∙ Uo ∙ A4 - αε^2 ∙ E ∙ I ∙ wo ∙ B4 + αε ∙ Mo ∙ C4 + Ho ∙ D4')
    text_abzac('где:')
    text_abzac('K = {} тс/м4  -  коэффициент пропорциональности;'.format(K))
    text_abzac('αε = {} (1/м)  -  коэффициент деформации;'.format(ae))
    text_abzac('E = {} тс/м2  -  модуль упругости материала сваи;'.format(E))
    text_abzac('I = {} м4  -  момент инерции поперечного сечения сваи, м4;'.format(I))
    text_abzac('Ho, Mo, Uo и ψo  -  то же, что и в Д.4 и Д.5;')
    
    def GMQ_otchet():
        global A1, B1, Ce1, D1, A3, B3, Ce3, D3, A4, B4, Ce4, D4, R, Gz, Mz, Qz, y1, fi1, C1
        GzMzQz(Z, Zp)
        text_abzac('Коэффициенты, значения которых принимают по таблице Д.3 СП 50-102-2003:')
        text_abzac('A1 = {}, B1 = {}, C1 = {}, D1 = {}'.format(A1, B1, Ce1, D1))
        text_abzac('A3 = {}, B3 = {}, C3 = {}, D3 = {}'.format(A3, B3, Ce3, D3))
        text_abzac('A4 = {}, B4 = {}, C4 = {}, D4 = {}'.format(A4, B4, Ce4, D4))
        ui.textEdit_2.append('')
        text_abzac_color('Gz = (({} тс/м4)/({} (1/м))) ∙ {} ∙ ({} м ∙ {} - {} ∙ {}/{} (1/м) + {} тсм ∙ {}/(({} (1/м))^2 ∙ {} тс/м2 ∙ {} м4) + {} тс ∙ {}/(({} (1/м))^3 ∙ {} тс/м2 ∙ {} м4)) = {} тс/м2'.format(K, ae, Zp, Uo, A1, wo, B1, ae, Mo, Ce1, ae, E, I, Ho, D1, ae, E, I, Gz))
        text_abzac_color('Mz = {}^2 (1/м) ∙ {} тс/м2 ∙ {} м4 ∙ {} м ∙ {} - {} (1/м) ∙ {} тс/м2 ∙ {} м4 ∙ {} рад ∙ {} + {} тсм ∙ {} + {} тс ∙ {}/{} (1/м) = {} тсм'.format(ae, E, I, Uo, A3, ae, E, I, wo, B3, Mo, Ce3, Ho, D3, ae, Mz))
        text_abzac_color('Qz = {}^3 (1/м) ∙ {} тс/м2 ∙ {} м4 ∙ {} м ∙ {} - {}^2 (1/м) ∙ {} тс/м2 ∙ {} м4 ∙ {} рад ∙ {} + {} (1/м) ∙ {} тсм ∙ {} + {} тс ∙ {} = {} тс'.format(ae, E, I, Uo, A4, ae, E, I, wo, B4, ae, Mo, Ce4, Ho, D4, Qz))
        text_abzac_color('Nz = N = {} тс'.format(Ns))

        text_abzac('В.8    Возможность использования линейных зависимостей при расчете свай должна проверяться по условию ограничения расчетного давления  , оказываемого на грунт боковыми поверхностями свай:')
        text_centr('Gz ≤ η1 ∙ η2 ∙ (4/cos(φI)) ∙ (γI ∙ Z ∙ tg(φI) + ξ ∙ CI)')
        text_abzac('где:')
        text_abzac('Gz = {} тс/м2  -  расчетное давление на грунт;'.format(Gz))
        text_abzac('Z = 0.85 / ae = 0.85 / {} = {} м  -  расчетное давление на грунт при lp > 2.5 м;'.format(ae, Z))
        text_abzac('γI = {} тс/м3  -  расчетный удельный (объемный) вес грунта ненарушенной структуры; '.format(y1))
        text_abzac('φI = {} рад  -  расчетные значения угла внутреннего трения грунта;'.format(fi1))
        text_abzac('СI = {} тс/м2  -  расчетные значения удельного сцепления грунта;'.format(C1))
        text_abzac('ξ = {}  -  коэффициент, принимаемый для забивных свай и свай-оболочек;'.format(ks))
        text_abzac('η1 = {}  -  коэффициент;'.format(h1))
        text_abzac('η2 = (Mc + Mt) / (n ∙ Mc + Mt) = ({} тсм + {} тсм) / ({} ∙ {} тсм + {} тсм) = {}  -  коэффициент, учитывающий долю постоянной нагрузки в суммарной нагрузке, где:'.format(Mc, Mt, n, Mc, Mt, h2))
        text_abzac('Mc = M ∙ (1 - fM) + H ∙ l1 ∙ (1 - fH) = {} тсм ∙ (1 - {}) + {} тс ∙ {} м ∙ (1 - {}) = {} тсм  -  момент от внешних постоянных нагрузок в сечении фундамента на уровне условной заделки на глубине l1 по п. 7.1.8;'.format(Ms, fM, Hs, l1, fH, Mc))
        text_abzac('Mt = M ∙ fM  + H ∙ l1 ∙ fH = {} тсм ∙ {} + {} тс ∙ {} м ∙ {} = {} тсм  -  то же, от внешних временных расчетных нагрузок;'.format(Ms, fM, Hs, l1, fH, Mt))
        text_abzac('n = {}  -  коэффициент для фундаментов с однорядным расположением свай на внецентренно приложенную вертикальную сжимающую нагрузку;'.format(n))
        text_abzac('fM = {}  -  отношение M от кратковременной нагрузки к M от полной нагрузки;'.format(fM))
        text_abzac('fH = {}  -  отношение H от кратковременной нагрузки к H от полной нагрузки;'.format(fH))
        ff = round(Gz / R * 100, 1)
        if Gz <= R:
            text_centr('{} тс/м2 ≤ {} ∙ {} ∙ (4/cos({})) ∙ ({} тс/м3 ∙ {} м ∙ tg({}) + {} ∙ {} тс/м2) = {} тс/м2'.format(Gz, h1, h2, fi1, y1, Z, fi1, ks, C1, R))
            text_centr('({} % от предельного значения) - условие выполнено '.format(ff))
            ui.textEdit_2.append('')
        else:
            ui.textEdit_2.setTextColor(QtGui.QColor (255, 0, 0))
            ui.textEdit_2.append('')
            text_centr_black('{} ≥ {} ∙ {} ∙ (4 / cos ({})) ∙ ({} ∙ {} ∙ tg({}) + {} ∙ {}) = {}'.format(Gz, h1, h2, fi1, y1, Z, fi1, ks, C1, R))
            text_centr_black('({} % от предельного значения) - условие НЕ выполнено '.format(ff))
            ui.textEdit_2.append('')
            ui.textEdit_2.setTextColor(QtGui.QColor (0, 0, 0))

    if lp > 2.5:
        Z = round(0.85 / ae, 4); Zp = round(Z * ae, 3)
        text_abzac('Zp = αε ∙ Z = {} (1/м) ∙ {} м = {}  -  приведенная глубина погружения сваи в грунт, где Z = 0.85/αε = {} м;'.format(ae, Z, Zp, Z))
        GMQ_otchet()
    else:
        Z = round(l / 3, 5); Zp = round(Z * ae, 5)
        ui.textEdit_2.append('')
        text_abzac_color('Т.к. lp ≤ 2,5 расчет проводим на двух глубинах, соответствующих Z = l/3  и Z = l')
        text_abzac_color('Расчет на глубине Z = l/3 = {} м/3 = {} м'.format(l, Z))
        text_abzac('Zp = αε ∙ Z = {} (1/м) ∙ {} м = {}  -  приведенная глубина погружения сваи в грунт;'.format(ae, Z, Zp))
        GMQ_otchet()
        Z = round(l, 5); Zp = round(Z * ae, 5)
        ui.textEdit_2.append('')
        text_abzac_color('Расчет на глубине Z = l = {} м'.format(Z))
        text_abzac('Zp = αε ∙ Z = {} (1/м) ∙ {} м = {}  -  приведенная глубина погружения сваи в грунт;'.format(ae, Z, Zp))
        GMQ_otchet()

# -----------------------------------------------------------------------------
# -----------------------------------------------------------------------------

def text_abzac_color_000(x):
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 100, 150))
    ui.textEdit_2.setFontWeight(100) # жирный текст
    ui.textEdit_2.append('{}'.format(x))
    ui.textEdit_2.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 0, 0))
    ui.textEdit_2.setFontWeight(1) # жирный текст
    # ui.textEdit_2.append('')

def text_abzac_color_111(x):
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 100, 150))
    ui.textEdit_2.setFontWeight(100) # жирный текст
    ui.textEdit_2.insertPlainText('{}'.format(x))
    ui.textEdit_2.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 0, 0))
    ui.textEdit_2.setFontWeight(1) # жирный текст

myDocumentsPath = QtCore.QDir.homePath() + "\Documents"
savePathFile = os.getcwd() + "\savePath.file"

"""Если файл с сохраненным путем существует и не пустой"""
if os.path.exists(savePathFile) == True and savePathFile != '':
    f = open(savePathFile, 'rb')        # чтение из файла
    loadPath = pickle.load(f)           # извлекаем ообъект из файла
    f.close()
else: 
    loadPath = myDocumentsPath

def writeFail(katalog):
    global loadPath
    '''Сохраняем путь к последнему сохараненному файлу в корневай папки программы'''
    f = open(savePathFile, 'wb')                # Запись в файл
    pickle.dump(katalog, f)                     # помещаем объект в файл
    f.close()

def soxranka():
    global loadPath
    ff = QtWidgets.QFileDialog.getSaveFileName(Form, 'Сохранить как', loadPath, 'NMH ( *.nmh )')
    katalog = ff [0]

    if '-nmh' not in katalog:
        katalog = katalog[:-4] + '-nmh' + katalog[-4:]

    if katalog != '':
        failName = os.path.split(katalog)[1]
        loadPath = writeFail(katalog)

        sbor_dannih()
        savex = [dannie, ige_skv, nni, ige_xap, Jl_xap, e_xap, γ1_xap, С1_xap, fi1_xap, tip_grunta_index]
        f = open(katalog, 'wb') # Запись в файл
        pickle.dump(savex, f) # помещаем объект в файл
        f.close()
        del savex # уничтожаем переменную savex  
        # ------------------------------------------------------------
        section = document.sections[0]
        section.top_margin  = Cm(2)     # верхнее поле
        section.bottom_margin = Cm(2)   # нижнее поле
        section.left_margin = Cm(3)     # левое поле
        section.right_margin = Cm(1.5)    # правое поле
        # ------------------------------------------------------------
        try:
            document.save('{}.docx'.format(katalog[:-4]))
        except:
            error_show(f"Файл {katalog} открыт. Закройте файл для перезаписи . . .")
            return
        Form.setWindowTitle(_translate("Form", "Pail - {}".format(failName)))
        ui.label_26.setText(_translate("Form", "Сохранение"))
        ui.textEdit_2.setText('')
        ui.textEdit_2.insertPlainText('Данные сохранены в файл: ')
        text_abzac_color_111(f'{failName}')
        ui.textEdit_2.append('')
        ui.textEdit_2.insertPlainText('Нажмите кнопку - ')
        text_abzac_color_111(f'Расчет')
    else: return None
    
def otkrivalka():
    global loadPath
    ff = QtWidgets.QFileDialog.getOpenFileName(Form, 'Открыть', loadPath, 'NMH ( *.nmh )')
    katalog = ff[0]
    
    if katalog != '':
        failName = os.path.split(katalog)[1]
        loadPath = writeFail(katalog)

        f = open(katalog, 'rb') # чтение из файла
        loadx = pickle.load(f) # извлекаем ообъект из файла
        for i in range(0, len(loadx)-1):
            for z in range(0, len(loadx[i])):
                if loadx[i][z] == 0:
                    loadx[i][z] = ''
        def vvod_open(a, b, c):
            _translate = QtCore.QCoreApplication.translate
            xxx = eval('[\'\' for z in range(0, ui.tableWidget{}.rowCount())]'.format(a))
            for i in range(0, len(xxx)):
                eval ('ui.tableWidget{}.item({}, {}).setText(_translate("Form", str({})))'.format(a, i, b, xxx[i]))
            for i in range(0, len(c)):
                # eval ('ui.tableWidget{}.item({}, {}).setText(_translate("Form", str({})))'.format(a, i, b, c[i]))
                eval ('ui.tableWidget{}.item({}, {}).setText(_translate("Form", "{}"))'.format(a, i, b, c[i]))
        vvod_open('_7', 0, loadx[0])
        vvod_open('_5', 0, loadx[1])
        vvod_open('_5', 1, loadx[2])
        vvod_open('_3', 0, loadx[3])
        vvod_open('_3', 1, loadx[4])
        vvod_open('_3', 2, loadx[5])
        vvod_open('_3', 3, loadx[6])
        vvod_open('_3', 4, loadx[7])
        vvod_open('_3', 5, loadx[8])
        for i in range(0, 10):
            eval('ui.comboBox_{}.setCurrentIndex({})'.format(i, loadx[9][i]))
        # ui.radioButton.setChecked(loadx[6])
        del loadx # уничтожаем переменную savex  
        ui.label_26.setText(_translate("Form", "Открытие"))
        ui.textEdit_2.setText('')
        ui.textEdit_2.insertPlainText('Данные закгужены из файла: ')
        text_abzac_color_111(f'{failName}')
        Form.setWindowTitle(_translate("Form", "Pail - {}".format(failName)))
        ui.textEdit_2.append('')
        ui.textEdit_2.insertPlainText('Нажмите кнопку - ')
        text_abzac_color_111(f'Расчет')
        return None
    else: return None
# ----------------------------------------------------------------------------------------
def text_centr_000(x):
    ui.textEdit_2.append('')
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 100, 150)) # цвет текста
    ui.textEdit_2.setFontWeight(100) # жирный текст
    ui.textEdit_2.append(x)
    ui.textEdit_2.setAlignment(QtCore.Qt.AlignHCenter) # центруем текст внутри абзаца
    ui.textEdit_2.setFontWeight(1) # убираем жирный текст
    ui.textEdit_2.setTextColor(QtGui.QColor (0, 0, 0))
def text_abzac_000(x):
    ui.textEdit_2.append('       {}'.format(x))
    ui.textEdit_2.setAlignment(QtCore.Qt.AlignLeft) # центруем текст внутри абзаца слева

def spravka():
    ui.label_26.setText(_translate("Form", "Справка"))
    ui.textEdit_2.setText('')
    text_abzac_color_000('       ИГЭ')
    ui.textEdit_2.insertPlainText(' – инженерно-геологический элемент')
    text_abzac_color_000('       γI')
    ui.textEdit_2.insertPlainText(' – удельный вес грунта')
    text_abzac_color_000('       фI')
    ui.textEdit_2.insertPlainText(' – угол внутреннего трения')
    text_abzac_color_000('       С1')
    ui.textEdit_2.insertPlainText(' – удельное сцепление грунта')
    text_abzac_color_000('       Jl')
    ui.textEdit_2.insertPlainText(' – показатель текучести грунта')
    text_abzac_color_000('       e')
    ui.textEdit_2.insertPlainText(' – коэффициент пористости')
    text_abzac_color_000('       H')
    ui.textEdit_2.insertPlainText(' – толщина i-го слоя грунта')
    text_abzac_color_000('       Uu')
    ui.textEdit_2.insertPlainText(' – предельные допустимые значения горизонтального перемещения головы сваи')
    text_abzac_color_000('       Ψu')
    ui.textEdit_2.insertPlainText(' – предельный угол поворота сваи')
    ui.textEdit_2.append('')
    text_abzac_000('При сохранении Расчета сохраняются два файла, один из них в формате *.docx')

# -----------------------------------------------------------------------------
# '''Авторизация'''
# noyblock = 'no'
# def block():
#     global noyblock
#     import datetime
#     dt_now = datetime.datetime.now().day
#     passeord = QtWidgets.QInputDialog.getText(Form, 'Авторизация', 'Введите пароль')
#     if str(dt_now) in passeord[0]:
#         noyblock = 'Yes'
#     else:
#         noyblock = 'no'
#         QtWidgets.QMessageBox.information(Form, 'Авторизация', 'Неверный пароль')

# def raschet_block():
#     global noyblock
#     if noyblock == 'no':
#         block()
#     if noyblock == 'Yes':
#         raschet()
# -----------------------------------------------------------------------------
# ui.pushButton_7.clicked.connect (raschet_block)
ui.pushButton_7.clicked.connect (raschet)
ui.pushButton_5.clicked.connect (soxranka)
ui.pushButton_4.clicked.connect (otkrivalka)
ui.pushButton_6.clicked.connect (spravka)
# -----------------------------------------------------------------------------
if __name__ == "__main__":
    sys.exit(app.exec_())
# -----------------------------------------------------------------------------
# ++++++++++++++++++++++++++++++++++++++++++++++++